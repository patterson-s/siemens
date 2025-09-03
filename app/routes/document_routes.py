import os
import tempfile
from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from app.models.models import Document, DocumentType, Project
from app import db
from cryptography.fernet import Fernet
import logging

bp = Blueprint('documents', __name__)
logger = logging.getLogger(__name__)

def allowed_file(filename):
    """Check if the file extension is allowed"""
    ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file, extraction_method='auto', use_ocr=False):
    """Extract text from various file types without requiring Tesseract"""
    filename = file.filename.lower()
    temp_path = os.path.join(tempfile.gettempdir(), secure_filename(file.filename))
    file.save(temp_path)
    
    extracted_text = ""
    methods_tried = []
    
    try:
        # Standard extraction for common file types
        if extraction_method in ['auto', 'standard']:
            methods_tried.append('standard')
            extracted_text = try_standard_extraction(temp_path, filename)
        
        # If standard method failed or robust was selected
        if (not extracted_text or len(extracted_text.split()) < 20) or extraction_method == 'robust':
            methods_tried.append('robust')
            extracted_text = try_robust_extraction(temp_path, filename)
        
        # If OCR was selected or previous methods failed
        if use_ocr or extraction_method == 'ocr_only' or (not extracted_text or len(extracted_text.split()) < 20):
            methods_tried.append('python_ocr')
            extracted_text = try_python_ocr(temp_path, filename)
        
        # Log extraction results
        word_count = len(extracted_text.split())
        logger.info(f"Extracted {word_count} words using methods: {', '.join(methods_tried)}")
        
        if word_count < 20:
            logger.warning(f"Low word count ({word_count}) for file: {filename}")
    
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        extracted_text = f"Error extracting text: {str(e)}"
    
    finally:
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    return extracted_text

def try_standard_extraction(file_path, filename):
    """Standard text extraction for common document types"""
    try:
        if filename.endswith('.pdf'):
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                return text
                
        elif filename.endswith('.docx'):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
            
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        return ""
    except Exception as e:
        logger.error(f"Standard extraction failed: {str(e)}")
        return ""

def try_robust_extraction(file_path, filename):
    """More robust extraction using specialized libraries"""
    try:
        # Try to use textract if available
        try:
            import textract
            return textract.process(file_path).decode('utf-8')
        except ImportError:
            logger.warning("textract not available, falling back to other methods")
        
        # If textract is not available, try other methods
        if filename.endswith('.pdf'):
            # Try multiple PDF libraries for better coverage
            
            # Try pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                    if text.strip():  # If we got meaningful text
                        return text
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
            
            # Try PyMuPDF (fitz)
            try:
                import fitz
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                if text.strip():  # If we got meaningful text
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        elif filename.endswith('.docx'):
            # Try a different approach for DOCX
            try:
                import docx2txt
                return docx2txt.process(file_path)
            except ImportError:
                logger.warning("docx2txt not available")
        
        return ""
    except Exception as e:
        logger.error(f"Robust extraction failed: {str(e)}")
        return ""

def try_python_ocr(file_path, filename):
    """Try OCR using PyMuPDF's built-in OCR capabilities"""
    try:
        # For PDFs, use PyMuPDF's text extraction which can handle some image-based text
        if filename.endswith('.pdf'):
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            
            # Try to extract text with different methods
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Method 1: Standard text extraction
                text += page.get_text() + "\n"
                
                # Method 2: Try to extract text from images on the page
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    # We could process this image with a simple OCR if available
                    # But for now, we'll just note that there are images
                    if not text.strip():
                        text += f"[Image-based content detected on page {page_num+1}]\n"
                        break
            
            return text
        
        # For image files, we can't do much without a proper OCR library
        # Just return a placeholder message
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            return "[Image file detected. OCR processing not available without additional libraries.]"
        
        return ""
    except Exception as e:
        logger.error(f"Python OCR alternative failed: {str(e)}")
        return ""

@bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_document():
    if request.method == 'GET':
        # Get all projects for the dropdown
        projects = Project.query.filter_by(active=True).order_by(Project.name).all()
        document_types = [choice for choice in DocumentType.choices()]
        return render_template('documents/upload.html', 
                             projects=projects,
                             document_types=document_types)
    
    # POST method - handle multiple file uploads
    project_id = request.form.get('project_id')
    if not project_id:
        flash('Please select a project', 'danger')
        return redirect(request.url)
    
    # Get advanced options
    extraction_method = request.form.get('extraction_method', 'auto')
    use_ocr = 'use_ocr' in request.form
    
    # Define document types and their corresponding file inputs
    document_types = {
        'external_evaluation_file': DocumentType.external_evaluation,
        'final_project_report_file': DocumentType.final_project_report,
        'original_proposal_file': DocumentType.original_proposal
    }
    
    uploaded_documents = []
    errors = []
    
    # Process each document type
    for file_input_name, document_type in document_types.items():
        if file_input_name in request.files:
            file = request.files[file_input_name]
            if file and file.filename != '' and allowed_file(file.filename):
                try:
                    # Extract content from the file with advanced options
                    content = extract_text_from_file(file, extraction_method, use_ocr)
                    
                    # Check if extraction was successful
                    word_count = len(content.split())
                    if word_count < 20:
                        errors.append(f'Warning: {file.filename} - Only {word_count} words were extracted. The document might be protected or primarily image-based.')
                    
                    # Create a preview
                    content_preview = content[:500] + "..." if len(content) > 500 else content
                    
                    # Create the document
                    document = Document(
                        project_id=project_id,
                        filename=secure_filename(file.filename),
                        file_type=file.content_type[:255] if hasattr(file, 'content_type') else file.filename.rsplit('.', 1)[1].lower(),
                        document_type=document_type.value,
                        file_size=word_count,  # Word count
                        content_preview=content_preview
                    )
                    
                    # Set the encrypted content
                    document.set_content(content)
                    
                    db.session.add(document)
                    uploaded_documents.append({
                        'filename': file.filename,
                        'type': document_type.value.replace('_', ' ').title(),
                        'word_count': word_count
                    })
                    
                except Exception as e:
                    logger.error(f"Error processing document {file.filename}: {str(e)}")
                    errors.append(f'Error processing {file.filename}: {str(e)}')
    
    # Commit all documents at once
    if uploaded_documents:
        try:
            db.session.commit()
            success_message = f'Successfully uploaded {len(uploaded_documents)} document(s): '
            for doc in uploaded_documents:
                success_message += f'{doc["filename"]} ({doc["word_count"]} words), '
            success_message = success_message.rstrip(', ')
            flash(success_message, 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Database error: {str(e)}', 'danger')
            return redirect(request.url)
    
    # Show any errors
    for error in errors:
        flash(error, 'warning')
    
    # If no documents were uploaded
    if not uploaded_documents and not errors:
        flash('No valid documents were selected for upload', 'warning')
    
    return redirect(request.url)

@bp.route('/view/<int:id>')
@login_required
def view(id):
    document = Document.query.get_or_404(id)
    # Add authorization check here if needed
    
    content = document.get_content()
    
    return render_template('documents/view.html', document=document, content=content)

@bp.route('/delete/<int:id>', methods=['POST'])
@login_required
def delete_document(id):
    document = Document.query.get_or_404(id)
    # Add authorization check here if needed
    
    project_id = document.project_id
    
    db.session.delete(document)
    db.session.commit()
    
    flash('Document deleted successfully', 'success')
    return redirect(url_for('main.project_details', project_id=project_id)) 