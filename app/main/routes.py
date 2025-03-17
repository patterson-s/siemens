from flask import render_template, redirect, url_for, request, flash, send_from_directory, jsonify, session, current_app, send_file
from app.main import bp
from app.models.models import Project, ProjectQuestion, Document, EvaluationRun, EvaluationResponse, EvaluationStatus, APILog, ChatSession, ChatMessage, DocumentType, Interview, InterviewQuestion  # Import InterviewQuestion
from app import db  # Import the db instance
import os
from anthropic import Anthropic, HUMAN_PROMPT, AI_PROMPT
from config import Config
from threading import Thread
import pdfplumber
import time
from datetime import datetime, timedelta  # Add timedelta to the import
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument
from io import BytesIO
from flask_login import login_required, current_user
from app.auth.decorators import admin_required
from sqlalchemy import desc, func, text, distinct, case, and_
from openai import OpenAI
import tiktoken  # Import tiktoken for token counting
from flask_wtf.csrf import generate_csrf
from cryptography.fernet import Fernet  # Add this at the top with other imports
import requests  # Import requests to make HTTP calls
from langchain.agents import create_sql_agent, AgentExecutor
from langchain.agents.agent_types import AgentType
from langchain.sql_database import SQLDatabase
from langchain_anthropic import ChatAnthropic
from langchain.agents.agent_toolkits import SQLDatabaseToolkit
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationChain
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits import SQLDatabaseToolkit
from langchain.agents.format_scratchpad import format_to_openai_functions
from langchain.agents.output_parsers import OpenAIFunctionsAgentOutputParser
import logging
import json  # Import the json module
import re  # Import re for regular expressions
import io
import csv

try:
    import xlsxwriter
except ImportError:
    pass  # Handle gracefully if xlsxwriter is not installed

# Configure logging to output to the console
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Update the DOCUMENT_TYPES tuple to show in the desired order
DOCUMENT_TYPES = [
    (DocumentType.original_proposal.value, 'Original Proposal'),
    (DocumentType.final_project_report.value, 'Project Report'),
    (DocumentType.external_evaluation.value, 'External Evaluation')
]

@bp.route('/')
@bp.route('/index')
@login_required
def index():
    # Get counts for dashboard
    total_projects = Project.query.count()
    total_documents = Document.query.count()
    
    # Get all projects for calculations
    all_projects = Project.query.all()
    
    # Count projects with at least one response
    projects_with_responses = db.session.query(Project.id).join(
        EvaluationResponse, Project.id == EvaluationResponse.project_id
    ).distinct().count()
    
    # Count projects with at least one document
    projects_with_docs = db.session.query(Project.id).join(
        Document, Project.id == Document.project_id
    ).distinct().count()
    
    # Count projects with ALL responses GENERATED
    projects_with_all_responses = get_projects_with_all_responses()
    
    # Count projects with ALL responses REVIEWED
    projects_with_all_reviewed = get_projects_with_all_reviewed()
    
    # Get recent projects
    recent_projects = Project.query.order_by(Project.id.desc()).limit(5).all()
    
    return render_template('main/index.html',
                          total_projects=total_projects,
                          total_documents=total_documents,
                          projects_with_responses=projects_with_responses,
                          projects_with_all_responses=projects_with_all_responses,
                          projects_with_all_reviewed=projects_with_all_reviewed,
                          projects_with_docs=projects_with_docs,
                          projects=all_projects,
                          recent_projects=recent_projects)

@bp.route('/projects')
@login_required
def projects():
    search_query = request.args.get('search', '')
    sort_by = request.args.get('sort', 'default')

    # Get total number of questions
    total_questions = ProjectQuestion.query.count()

    # Subquery to get the most recent response for each question per project
    latest_responses = db.session.query(
        EvaluationResponse.project_id,
        EvaluationResponse.question_id,
        func.max(EvaluationResponse.created_at).label('max_created_at')
    ).group_by(
        EvaluationResponse.project_id,
        EvaluationResponse.question_id
    ).subquery()

    # Query to get response counts and review counts for each project
    response_stats = db.session.query(
        EvaluationResponse.project_id,
        func.count(distinct(EvaluationResponse.question_id)).label('answered_questions'),
        func.sum(case((EvaluationResponse.reviewed == True, 1), else_=0)).label('reviewed_count')
    ).join(
        latest_responses,
        and_(
            EvaluationResponse.project_id == latest_responses.c.project_id,
            EvaluationResponse.question_id == latest_responses.c.question_id,
            EvaluationResponse.created_at == latest_responses.c.max_created_at
        )
    ).group_by(
        EvaluationResponse.project_id
    ).subquery()

    # Query only active projects from the database
    active_projects = Project.query.filter_by(active=True).outerjoin(
        response_stats,
        Project.id == response_stats.c.project_id
    )

    # Apply search filter
    if search_query:
        active_projects = active_projects.filter(Project.integrity_partner_name.ilike(f'%{search_query}%'))

    # Apply sorting
    if sort_by == 'default' or sort_by == 'funding_round':
        active_projects = active_projects.order_by(Project.name_of_round.asc(), Project.file_number_db.asc())
    elif sort_by == 'file_number':
        active_projects = active_projects.order_by(Project.file_number_db.asc())
    elif sort_by == 'partner_name':
        active_projects = active_projects.order_by(Project.integrity_partner_name.asc())
    elif sort_by == 'region':
        active_projects = active_projects.order_by(Project.region.asc())
    elif sort_by == 'countries':
        active_projects = active_projects.order_by(Project.countries_covered.asc())

    active_projects = active_projects.all()

    # Convert query results to dictionary with tuple values
    stats_dict = {}
    for project_id, answered, reviewed in db.session.query(
        response_stats.c.project_id,
        response_stats.c.answered_questions,
        response_stats.c.reviewed_count
    ).all():
        stats_dict[project_id] = (answered or 0, reviewed or 0)

    return render_template('main/projects.html', 
                         projects=active_projects,
                         total_questions=total_questions,
                         response_counts=stats_dict,
                         title='Projects - ' + Config.APP_NAME)

@bp.route('/questions')
@login_required
def questions():
    # Query all project questions from the database
    all_questions = ProjectQuestion.query.all()
    return render_template('main/questions.html', questions=all_questions, title='Questions - ' + Config.APP_NAME)

@bp.route('/questions/edit/<int:question_id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_question(question_id):
    question = ProjectQuestion.query.get_or_404(question_id)
    
    if request.method == 'POST':
        try:
            question.question = request.form['question']
            question.name = request.form['name']
            question.prompt = request.form['prompt']
            question.order = int(request.form['order'])  # Convert to integer
            db.session.commit()
            flash('Project question updated successfully!', 'success')
            return redirect(url_for('main.questions'))
        except ValueError:
            flash('Invalid order value. Please enter a number.', 'danger')
            return redirect(request.url)
        except Exception as e:
            flash(f'Error updating question: {str(e)}', 'danger')
            return redirect(request.url)

    return render_template('main/edit_question.html', question=question)

@bp.route('/projects/<int:project_id>')
@login_required
def project_details(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Get the most recent response for each question using project_id
    latest_responses = db.session.query(
        EvaluationResponse
    ).filter(
        EvaluationResponse.project_id == project_id
    ).distinct(
        EvaluationResponse.question_id
    ).order_by(
        EvaluationResponse.question_id,
        EvaluationResponse.created_at.desc()
    ).all()

    return render_template('main/project_details.html', 
                         project=project,
                         latest_responses=latest_responses,
                         csrf_token=generate_csrf())

def extract_text_from_file(file):
    """Extract text content from uploaded files."""
    # Get file extension
    file_ext = file.filename.rsplit('.', 1)[1].lower()
    
    # Read file content
    file_content = file.read()
    
    # Reset file pointer for content extraction
    file.seek(0)
    
    # Extract text based on file type
    if file_ext == 'pdf':
        return extract_text_from_pdf(file)
    elif file_ext == 'txt':
        return file_content.decode('utf-8')
    elif file_ext in ['doc', 'docx']:
        return extract_text_from_doc(file)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

@bp.route('/projects/<int:project_id>/upload', methods=['GET', 'POST'])
@login_required
def upload_document(project_id):
    project = Project.query.get_or_404(project_id)
    
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file selected', 'danger')
            return redirect(request.url)
            
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'danger')
            return redirect(request.url)
            
        if file and allowed_file(file.filename):
            try:
                # Extract text from the document
                text = extract_text_from_file(file)
                
                # Calculate word count
                word_count = len(text.split())
                
                # Create new document
                document = Document(
                    project_id=project.id,
                    filename=secure_filename(file.filename),
                    file_type=file.content_type,
                    document_type=request.form.get('document_type'),
                    file_size=word_count,  # Store word count instead of file size
                    content_preview=text[:1000] if text else None
                )
                
                # Encrypt and set the content
                try:
                    document.set_content(text)
                    print(f"Content encrypted successfully")  # Debug log
                except Exception as e:
                    raise ValueError(f"Encryption failed: {str(e)}")
                
                db.session.add(document)
                db.session.commit()
                
                # Verify content was saved
                saved_doc = Document.query.get(document.id)
                if not saved_doc._content:
                    raise ValueError("Content was not saved to database")
                
                flash(f'Document uploaded successfully. Word count: {word_count}', 'success')
                return redirect(url_for('main.project_details', project_id=project.id))
                
            except Exception as e:
                db.session.rollback()
                flash(f'Error processing document: {str(e)}', 'danger')
                return redirect(request.url)
                
        else:
            flash('Invalid file type. Supported formats: PDF, TXT, DOC, DOCX', 'danger')
            return redirect(request.url)
    
    return render_template('main/upload_document.html', 
                         project=project,
                         document_types=DOCUMENT_TYPES)

@bp.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory('uploads', filename) 

@bp.route('/projects/<int:project_id>/documents/<int:document_id>/delete', methods=['POST'])
@login_required
def delete_document(project_id, document_id):
    try:
        from sqlalchemy import text
        
        # First check if document exists and belongs to this project
        document = Document.query.filter_by(id=document_id, project_id=project_id).first_or_404()
        
        # Remove any references in evaluation_documents
        db.session.execute(
            text('DELETE FROM univ_evaluation_documents WHERE document_id = :doc_id'),
            {'doc_id': document_id}
        )
        
        # Delete the document
        db.session.delete(document)
        db.session.commit()
        
        flash('Document deleted successfully.', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting document: {str(e)}', 'danger')
        print(f"Error deleting document: {str(e)}")
    
    return redirect(url_for('main.project_details', project_id=project_id))

@bp.route('/api/projects/<int:project_id>/process_document/<int:document_id>', methods=['POST'])
def process_document(project_id, document_id):
    document = Document.query.get_or_404(document_id)
    file_path = os.path.join('uploads', document.filename)
    print(f"File path: {file_path}")

    try:
        if document.file_type == 'application/pdf':
            # Handle PDF files
            with pdfplumber.open(file_path) as pdf:
                content = ""
                for page in pdf.pages:
                    content += page.extract_text() + "\n"
        else:
            # Handle text files
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()

        print(f"File content length: {len(content)}")
        document.content = content
        document.content_preview = content[:100]
        db.session.commit()
        return jsonify({'message': 'Document processed successfully!'}), 200

    except Exception as e:
        print(f"Error processing file: {e}")
        return jsonify({'error': str(e)}), 500

def count_tokens(text):
    """Count tokens using tiktoken with cl100k_base encoding (used by Claude)"""
    encoding = tiktoken.get_encoding("cl100k_base")
    return len(encoding.encode(text))

@bp.route('/projects/<int:project_id>/start_assessment', methods=['GET', 'POST'])
@login_required
def start_assessment(project_id):
    project = Project.query.get_or_404(project_id)
    questions = ProjectQuestion.query.order_by(ProjectQuestion.order).all()
    
    if request.method == 'POST':
        selected_docs = request.form.getlist('selected_documents')
        selected_questions = request.form.getlist('selected_questions')
        
        if not selected_docs or not selected_questions:
            flash('Please select at least one document and one question', 'warning')
            return redirect(request.url)
            
        try:
            # Get the selected documents' content
            documents = Document.query.filter(Document.id.in_(selected_docs)).all()
            total_doc_content = "\n\n".join(doc.get_content() for doc in documents)
            
            # Create new evaluation run
            evaluation = EvaluationRun(
                project_id=project.id,
                model_used=request.form.get('model_select', Config.CLAUDE_MODEL),
                system_prompt=Config.DEFAULT_SYSTEM_PROMPT,
                status=EvaluationStatus.IN_PROGRESS,
                total_questions=len(selected_questions)
            )
            
            # Add documents to evaluation
            evaluation.documents.extend(documents)
            
            # Calculate tokens for documents and system prompt once
            doc_and_system_tokens = count_tokens(f"{Config.DEFAULT_SYSTEM_PROMPT}\n\n{total_doc_content}")
            
            # Add tokens for each question prompt
            question_tokens = sum(
                count_tokens(question.prompt)
                for question in ProjectQuestion.query.filter(ProjectQuestion.id.in_(selected_questions)).all()
            )
            
            # Total input tokens is system prompt + documents + (questions * number of questions)
            evaluation.estimated_input_tokens = doc_and_system_tokens + question_tokens
            
            db.session.add(evaluation)
            db.session.commit()
            
            # Start the evaluation process in a background thread
            temperature = float(request.form.get('temperature', 0.5))
            thread = Thread(target=run_evaluation, 
                          args=(project.id, evaluation.id, selected_questions, temperature, selected_docs))
            thread.daemon = True
            thread.start()
            
            # Redirect to the assessment progress page
            return render_template('main/assessment_progress.html',
                                project=project,
                                total_questions=len(selected_questions),
                                evaluation_id=evaluation.id,
                                evaluation=evaluation)
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error starting assessment: {str(e)}', 'danger')
            return redirect(request.url)
    
    return render_template('main/start_assessment.html',
                         project=project,
                         questions=questions,
                         document_types=DOCUMENT_TYPES)

@bp.route('/projects/<int:project_id>/assessment/<int:evaluation_id>')
@login_required
def view_assessment(project_id, evaluation_id):
    evaluation = EvaluationRun.query.get_or_404(evaluation_id)
    # When passing responses to template, no changes needed as template will use get_response_text()
    return render_template('main/view_assessment.html', 
                         evaluation=evaluation,
                         project=evaluation.project)

@bp.route('/setup/add_test_questions')
def add_test_questions():
    # Check if questions already exist
    if ProjectQuestion.query.first() is not None:
        flash('Test questions already exist!', 'info')
        return redirect(url_for('main.questions'))
    
    # Add test questions
    test_questions = [
        {
            'name': 'Overall Impact',
            'question': 'What was the overall impact of the project?',
            'prompt': 'Based on the provided documentation, analyze and evaluate the overall impact of this project. Consider both intended and unintended outcomes.',
            'order': 1
        },
        {
            'name': 'Effectiveness',
            'question': 'How effective was the project in achieving its objectives?',
            'prompt': 'Evaluate the effectiveness of the project in achieving its stated objectives. Provide specific examples from the documentation.',
            'order': 2
        },
        {
            'name': 'Sustainability',
            'question': 'How sustainable are the project outcomes?',
            'prompt': 'Assess the sustainability of the project outcomes. Consider factors such as institutional capacity, financial sustainability, and long-term viability.',
            'order': 3
        },
        {
            'name': 'Lessons Learned',
            'question': 'What are the key lessons learned from this project?',
            'prompt': 'Identify and analyze the key lessons learned from this project. What worked well, what could have been done differently, and what recommendations can be made for similar projects?',
            'order': 4
        }
    ]
    
    for q in test_questions:
        question = ProjectQuestion(**q)
        db.session.add(question)
    
    db.session.commit()
    flash('Test questions added successfully!', 'success')
    return redirect(url_for('main.questions')) 

@bp.route('/projects/assessment_progress/<int:evaluation_id>')
@login_required
def check_assessment_progress(evaluation_id):
    evaluation = EvaluationRun.query.get_or_404(evaluation_id)
    
    # Count completed responses using correct column name
    completed_responses = EvaluationResponse.query.filter_by(evaluation_run_id=evaluation_id).count()
    
    # Get total output tokens from API logs
    output_tokens = db.session.query(
        func.sum(APILog.output_tokens)
    ).filter(
        APILog.evaluation_id == evaluation_id,
        APILog.success == True
    ).scalar() or 0
    
    # Calculate progress
    progress = {
        'status': evaluation.status.value,
        'total_questions': evaluation.total_questions,
        'questions_answered': completed_responses,
        'model_used': evaluation.model_used,
        'output_tokens': output_tokens,
        'is_completed': evaluation.status == EvaluationStatus.COMPLETED
    }
    
    return jsonify(progress)

def run_evaluation(project_id, evaluation_id, selected_questions, temperature, selected_docs):
    """Run the evaluation process in the background"""
    from app import create_app
    app = create_app()
    
    with app.app_context():
        try:
            evaluation = EvaluationRun.query.get(evaluation_id)
            project = Project.query.get(project_id)
            
            # Get selected documents with their types
            documents = Document.query.filter(Document.id.in_(selected_docs)).all()
            
            # Build context from documents with caching instruction
            document_contexts = []
            for doc in documents:
                # Fix: Convert the Enum to string before using replace()
                if hasattr(doc.document_type, 'value'):
                    # If it's an Enum with a value attribute
                    doc_type = str(doc.document_type.value)
                elif hasattr(doc.document_type, 'name'):
                    # If it's an Enum with a name attribute
                    doc_type = str(doc.document_type.name)
                else:
                    # Fallback to string conversion
                    doc_type = str(doc.document_type)
                
                formatted_type = doc_type.replace('_', ' ').title()
                
                doc_context = f"\n=== {formatted_type}: {doc.filename} ===\n{doc.get_content()}\n"
                document_contexts.append(doc_context)
            
            if not document_contexts:
                raise ValueError("No readable document content found")

            # Combine all document contexts with instruction
            file_instruction = ("The following content is from project documents. "
                              "Use this information when responding to questions. "
                              "Reference specific documents when citing information:\n\n")
            combined_context = file_instruction + "\n".join(document_contexts)
            
            # Initialize Anthropic client with caching headers
            client = Anthropic(
                api_key=Config.CLAUDE_API_KEY,
                default_headers={
                    "anthropic-version": "2023-06-01",
                    "anthropic-beta": "prompt-caching-2024-07-31"
                }
            )
            
            # Create the system content with caching
            system = [
                {
                    "type": "text",
                    "text": Config.DEFAULT_SYSTEM_PROMPT
                },
                {
                    "type": "text",
                    "text": combined_context,
                    "cache_control": {"type": "ephemeral"}
                }
            ]
            
            # Process each question
            completed_questions = 0
            for question in ProjectQuestion.query.filter(ProjectQuestion.id.in_(selected_questions)).order_by(ProjectQuestion.order):
                try:
                    # Create API log entry
                    api_log = APILog(
                        project_id=project.id,
                        evaluation_id=evaluation.id,
                        question_id=question.id,
                        model_used=evaluation.model_used,
                        project_name=project.name,
                        question_name=question.name,
                        start_time=datetime.utcnow()
                    )
                    db.session.add(api_log)
                    db.session.commit()
                    
                    # Call Claude API
                    response = client.messages.create(
                        model=evaluation.model_used,
                        temperature=temperature,
                        max_tokens=4000,
                        system=system,
                        messages=[{
                            "role": "user",
                            "content": f"""
Project Context:
Project Name: {project.name}
Project Objectives: {project.key_project_objectives}

Question: {question.question}
Detailed Instructions: {question.prompt}

Please provide your analysis based on the available documents."""
                        }]
                    )
                    
                    # Update API log with token usage and cost
                    total_tokens = response.usage.input_tokens + response.usage.output_tokens
                    api_log.input_tokens = response.usage.input_tokens
                    api_log.output_tokens = response.usage.output_tokens
                    api_log.cache_tokens = len(combined_context.split())  # Approximate token count for cached content
                    
                    # Calculate cost
                    model_costs = Config.CLAUDE_COSTS.get(evaluation.model_used, Config.CLAUDE_COSTS['claude-3-haiku-20240307'])
                    input_cost = (api_log.input_tokens / 1000) * model_costs['input']
                    output_cost = (api_log.output_tokens / 1000) * model_costs['output']
                    cache_cost = (api_log.cache_tokens / 1000) * model_costs['cache']
                    api_log.cost_usd = input_cost + output_cost + cache_cost
                    
                    api_log.success = True
                    api_log.end_time = datetime.utcnow()
                    
                    # Create evaluation response
                    evaluation_response = EvaluationResponse(
                        evaluation_run_id=evaluation.id,
                        question_id=question.id,
                        project_id=project_id,
                        reviewed=False,
                        created_at=datetime.utcnow()
                    )
                    # Use the property setter for encryption
                    evaluation_response.response_text = response.content[0].text
                    db.session.add(evaluation_response)
                    db.session.commit()
                    
                    # Adaptive delay based on token usage
                    if total_tokens > 30000:
                        time.sleep(5)
                    else:
                        time.sleep(2)
                    
                    # Update progress
                    completed_questions += 1
                    evaluation.completed_questions = completed_questions
                    evaluation.progress = int((completed_questions / len(selected_questions)) * 100)
                    db.session.commit()
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error processing question {question.id}: {error_msg}")
                    
                    api_log.success = False
                    api_log.error_message = error_msg
                    api_log.end_time = datetime.utcnow()
                    db.session.commit()
                    
                    if "rate_limit_error" in error_msg.lower():
                        print("Rate limit hit, waiting 60 seconds...")
                        time.sleep(60)
                        continue
                    else:
                        evaluation.status = EvaluationStatus.FAILED
                        db.session.commit()
                        raise e
            
            evaluation.status = EvaluationStatus.COMPLETED
            db.session.commit()
            
        except Exception as e:
            print(f"Evaluation failed: {str(e)}")
            evaluation = EvaluationRun.query.get(evaluation_id)
            if evaluation:
                evaluation.status = EvaluationStatus.FAILED
                db.session.commit()

def get_recent_tokens():
    """Get total tokens used in the last minute"""
    one_minute_ago = datetime.utcnow() - timedelta(minutes=1)
    return db.session.query(
        db.func.sum(APILog.input_tokens)
    ).filter(
        APILog.start_time >= one_minute_ago,
        APILog.success == True
    ).scalar() or 0

@bp.route('/projects/<int:project_id>/delete_assessment/<int:evaluation_id>', methods=['POST'])
@login_required
@admin_required
def delete_assessment(project_id, evaluation_id):
    evaluation = EvaluationRun.query.get_or_404(evaluation_id)
    
    # Only delete evaluation responses
    EvaluationResponse.query.filter_by(evaluation_run_id=evaluation.id).delete()
    
    # Delete the evaluation itself
    db.session.delete(evaluation)
    db.session.commit()
    
    flash('Assessment deleted successfully!', 'success')
    return redirect(url_for('main.project_details', project_id=project_id)) 

@bp.route('/projects/toggle_status/<int:project_id>', methods=['POST'])
@login_required
@admin_required
def toggle_project_status(project_id):
    project = Project.query.get_or_404(project_id)
    project.active = not project.active
    db.session.commit()
    status = "activated" if project.active else "deactivated"
    flash(f'Project {status} successfully!', 'success')
    return redirect(url_for('main.projects')) 

@bp.route('/admin/projects')
@login_required
@admin_required
def admin_projects():
    # Query all projects, including inactive ones
    all_projects = Project.query.order_by(Project.name).all()
    return render_template('main/admin_projects.html', 
                         projects=all_projects, 
                         title='Manage Projects - ' + Config.APP_NAME) 

@bp.route('/admin/api-logs')
@login_required
@admin_required
def api_logs():
    # Get summary statistics
    total_calls = APILog.query.count()
    successful_calls = APILog.query.filter_by(success=True).count()
    success_rate = (successful_calls / total_calls * 100) if total_calls > 0 else 0
    
    # Get average token usage for successful calls
    token_stats = db.session.query(
        db.func.avg(APILog.input_tokens).label('avg_input'),
        db.func.avg(APILog.output_tokens).label('avg_output')
    ).filter(APILog.success == True).first()
    
    # Get average response time by model
    model_stats = db.session.query(
        APILog.model_used,
        db.func.count(APILog.id).label('calls'),
        db.func.avg(APILog.end_time - APILog.start_time).label('avg_time'),
        db.func.avg(APILog.input_tokens).label('avg_input'),
        db.func.avg(APILog.output_tokens).label('avg_output'),
        db.func.sum(db.case((APILog.success == True, 1), else_=0)).label('successes')
    ).filter(
        APILog.end_time.isnot(None)
    ).group_by(
        APILog.model_used
    ).all()
    
    # Format model stats for display
    model_performance = [{
        'model': model.replace('claude-', 'Claude ').replace('-latest', '').replace('-', '.'),
        'calls': calls,
        'avg_time': round(avg_time.total_seconds(), 2) if avg_time else 0,
        'avg_input': int(avg_input or 0),
        'avg_output': int(avg_output or 0),
        'success_rate': (successes / calls * 100) if calls > 0 else 0
    } for model, calls, avg_time, avg_input, avg_output, successes in model_stats]
    
    # Get overall average response time
    avg_response_time = db.session.query(
        db.func.avg(APILog.end_time - APILog.start_time)
    ).filter(APILog.success == True).scalar() or timedelta(seconds=0)
    
    # Get detailed logs with related data
    logs = db.session.query(APILog, Project, ProjectQuestion)\
        .join(Project, APILog.project_id == Project.id)\
        .join(ProjectQuestion, APILog.question_id == ProjectQuestion.id)\
        .order_by(APILog.start_time.desc())\
        .all()
    
    return render_template('main/api_logs.html',
                         logs=logs,
                         total_calls=total_calls,
                         success_rate=success_rate,
                         avg_input_tokens=token_stats.avg_input,
                         avg_output_tokens=token_stats.avg_output,
                         avg_response_time=avg_response_time,
                         model_performance=model_performance,
                         title='API Logs - ' + Config.APP_NAME) 

@bp.route('/ai-agent', methods=['GET', 'POST'])
@login_required
def ai_agent():
    # Get or create chat session
    session_id = session.get('chat_session_id')
    if session_id:
        chat_session = ChatSession.query.get(session_id)
    else:
        chat_session = ChatSession()
        db.session.add(chat_session)
        db.session.commit()
        session['chat_session_id'] = chat_session.id

    if request.method == 'POST':
        user_query = request.form.get('query')
        
        # Create user message
        user_message = ChatMessage(
            session_id=chat_session.id,
            role='user',
            content=user_query
        )
        db.session.add(user_message)

        # Create SQLDatabase instance
        db_uri = current_app.config['SQLALCHEMY_DATABASE_URI']
        sql_db = SQLDatabase.from_uri(db_uri)

        # Initialize the LLM with higher temperature
        llm = ChatAnthropic(
            model="claude-3-5-sonnet-20241022",
            anthropic_api_key=Config.CLAUDE_API_KEY,
            temperature=0.3,  
            max_tokens=4000
        )

        # Create the toolkit with both db and llm
        toolkit = SQLDatabaseToolkit(db=sql_db, llm=llm)

        # Create the agent with the toolkit
        agent = create_sql_agent(llm, toolkit, agent_type=AgentType.ZERO_SHOT_REACT_DESCRIPTION)

        # Define the schema description
        schema_description = """
        The primary objects of interest are projects, question responses, and documents.
        
        The database contains the following tables:

        1. **univ_projects** - Projects
            - **id**: Integer, primary key - Unique identifier for each project.
            - **name_of_round**: Numeric(3, 1), nullable=False - The round name, allowing decimal values (e.g., 1.0, 2.5).
            - **name**: String(255), nullable=False - The name of the project.
            - **file_number_db**: Float - The file number associated with the project.
            - **scope**: String(100) - The scope of the project.
            - **region**: String(100) - The geographical region of the project.
            - **countries_covered**: String(255) - The countries covered by the project.
            - **integrity_partner_name**: String(255) - The name of the integrity partner involved in the project.
            - **partner_type**: String(50) - The type of partner involved in the project.
            - **project_partners**: Text - Details about the project partners.
            - **wb_or_eib**: String(50) - Indicates whether the project is associated with the World Bank (WB) or the European Investment Bank (EIB).
            - **key_project_objectives**: Text - The key objectives of the project.
            - **sectoral_scope**: String(100) - The sectoral scope of the project.
            - **specific_sector**: String(100) - The specific sector the project targets.
            - **funding_amount_usd**: Numeric(10, 2) - The funding amount in USD millions, allowing for two decimal places.
            - **duration**: String(50) - The duration of the project.
            - **start_year**: Integer - The year the project starts.
            - **end_year**: Integer - The year the project ends.
            - **wb_income_classification**: String(50) - The income classification according to the World Bank.
            - **corruption_quintile**: String(25) - The corruption quintile classification.
            - **cci**: Numeric(4, 2) - The corruption control index, allowing for two decimal places.
            - **government_type_eiu**: String(100) - The type of government according to the Economist Intelligence Unit (EIU).
            - **government_score_eiu**: Numeric(4, 2) - The government score according to the EIU, allowing for two decimal places.
            - **notes**: Text - Additional information about the project.


        2. **univ_documents** - Documents
            - **id**: Integer, primary key - Unique identifier for each document.
            - **project_id**: Integer, foreign key referencing univ_projects(id), nullable=False - The ID of the project associated with this document.
            - **filename**: String(255), nullable=False - The name of the document file.
            - **file_type**: String(100) - The type of the file (e.g., PDF, DOCX).
            - **document_type**: Enum(DocumentType), nullable=False, default=DocumentType.external_evaluation - The type of document, with a default value of 'external_evaluation'.
            - **file_size**: Integer - The word count of the content from the document (stored in content column).
            -  **content_preview**: Text - A preview of the document's content.
            - **content**: Text - The full content of the document.
            - **created_at**: DateTime, default=datetime.utcnow - The timestamp when the document was created.


        3. **univ_evaluation_responses** - Question Responses
            - **id**: Integer, primary key - Unique identifier for each response.
            - **question_id**: Integer, foreign key referencing univ_project_questions(id) - The ID of the question associated with this response.
            - **project_id**: Integer, foreign key referencing univ_projects(id) - The ID of the project associated with this response.
            - **evaluation_run_id**: Integer, foreign key referencing univ_evaluation_runs(id) - The ID of the evaluation run associated with this response.
            - **response_text**: Text, nullable=False - The text of the question response. **This field is encrypted. To access its content, you must use the decryption function.**
            - **reviewed**: Boolean, default=False - Indicates whether the response has been reviewed.
            - **created_at**: DateTime, default=datetime.utcnow - The timestamp when the response was created.

            **When retrieving responses, ensure you only retrieve the most recent response that has been generated for each project and question combination.**

        4. **univ_project_questions** - Questions   
            - **id**: Integer, primary key - Unique identifier for each question.
            - **question**: Text, nullable=False - The text of the question.
            - **name**: String(50), nullable=False - The short title of the question.
            - **prompt**: Text, nullable=False - The prompt used to generate the response.
            - **order**: Integer, nullable=False - The order of the question.
            - **created_at**: DateTime, default=datetime.utcnow - The timestamp when the question was created.

        **Important Note**: The 'response_text' field in the 'univ_evaluation_responses' table is encrypted. You cannot directly query or filter on this field. Instead, I will provide you with decrypted responses for your analysis.
        """

        try:
            # Extract project information from the user query
            project_info = None
            project_id = None
            
            # Check if the query mentions a specific project
            if "project" in user_query.lower():
                # Try to extract project name or ID
                project_match = re.search(r'project\s+(?:named|called|titled)?\s*["\']?([^"\']+)["\']?', user_query.lower())
                if project_match:
                    project_name = project_match.group(1).strip()
                    project = Project.query.filter(Project.name.ilike(f"%{project_name}%")).first()
                    if project:
                        project_id = project.id
                        project_info = f"Project ID: {project.id}, Name: {project.name}"
            
            # Retrieve and decrypt evaluation responses
            decrypted_responses = []
            
            if project_id:
                # If a specific project is mentioned, retrieve responses for that project
                evaluation_responses = EvaluationResponse.query.filter_by(project_id=project_id).all()
            else:
                # Otherwise, retrieve a limited number of responses
                evaluation_responses = EvaluationResponse.query.limit(50).all()
            
            # Decrypt the responses
            for response in evaluation_responses:
                decrypted_text = response.get_response_text()
                if decrypted_text and not decrypted_text.startswith("Error:"):
                    # Get the question text
                    question = ProjectQuestion.query.get(response.question_id)
                    question_text = question.question if question else "Unknown Question"
                    
                    # Get the project name
                    project = Project.query.get(response.project_id)
                    project_name = project.name if project else "Unknown Project"
                    
                    decrypted_responses.append({
                        "response_id": response.id,
                        "project_id": response.project_id,
                        "project_name": project_name,
                        "question_id": response.question_id,
                        "question_text": question_text,
                        "decrypted_content": decrypted_text
                    })
            
            # Format the decrypted responses for the AI agent
            formatted_responses = ""
            for resp in decrypted_responses:
                formatted_responses += f"\n--- Response ID: {resp['response_id']} ---\n"
                formatted_responses += f"Project: {resp['project_name']} (ID: {resp['project_id']})\n"
                formatted_responses += f"Question: {resp['question_text']} (ID: {resp['question_id']})\n"
                formatted_responses += f"Content: {resp['decrypted_content']}\n"
            
            # Combine schema description with user query and decrypted responses
            full_query = f"{schema_description}\n\nDecrypted Evaluation Responses:\n{formatted_responses}\n\nUser Query: {user_query}"
            
            # Invoke the agent with the full query
            try:
                # First try with handle_parsing_errors=True
                response = agent.invoke(full_query, handle_parsing_errors=True)
                
                # Ensure response is a string
                if isinstance(response, dict):
                    output = response.get("output", "No output available.")
                else:
                    output = str(response)
                    
            except Exception as parsing_error:
                # If there's still an error, extract the raw output from the error message
                error_str = str(parsing_error)
                if "Could not parse LLM output: `" in error_str:
                    # Extract the raw output from the error message
                    raw_output = error_str.split("Could not parse LLM output: `")[1].rsplit("`", 1)[0]
                    output = raw_output
                else:
                    # If we can't extract the output, use the error message
                    output = f"Error processing query: {error_str}"
                
                logger.warning(f"Parsing error handled: {error_str}")
                
        except Exception as e:  # Catch all exceptions
            logger.error("Error processing user query: %s", e)  # Log the error
            logger.debug("User query: %s", user_query)  # Log the user query for debugging
            db.session.rollback()  # Rollback the session to clear the transaction state
            output = "I'm sorry, but I couldn't process your request. Please try again or rephrase your question."

        # Create assistant message
        assistant_message = ChatMessage(
            session_id=chat_session.id,
            role='assistant',
            content=output  # Store the output as a string
        )
        db.session.add(assistant_message)
        db.session.commit()

        # Get all messages for this session
        messages = ChatMessage.query.filter_by(session_id=chat_session.id)\
            .order_by(ChatMessage.created_at.desc())\
            .all()

        return render_template('main/ai_agent.html',
                               response=output,  # Pass the output to the template
                               chat_history=messages,
                               current_session=chat_session,
                               all_sessions=ChatSession.query.order_by(ChatSession.created_at.desc()).all())

    # GET request handling
    messages = ChatMessage.query.filter_by(session_id=chat_session.id)\
        .order_by(ChatMessage.created_at.desc())\
        .all()

    return render_template('main/ai_agent.html',
                           response=None,
                           chat_history=messages,
                           current_session=chat_session,
                           all_sessions=ChatSession.query.order_by(ChatSession.created_at.desc()).all())

@bp.route('/ai-agent/new-session')
def new_chat_session():
    session.pop('chat_session_id', None)
    return redirect(url_for('main.ai_agent'))

@bp.route('/ai-agent/session/<int:session_id>')
def load_chat_session(session_id):
    chat_session = ChatSession.query.get_or_404(session_id)
    session['chat_session_id'] = chat_session.id
    return redirect(url_for('main.ai_agent'))

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx'}
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS 

def extract_text_from_pdf(file):
    """Extract text from a PDF file"""
    try:
        with pdfplumber.open(file) as pdf:
            text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                    
                # Also extract text from tables
                for table in page.extract_tables():
                    table_text = []
                    for row in table:
                        row_text = [cell.strip() if cell else '' for cell in row]
                        if any(row_text):  # Only add non-empty rows
                            table_text.append(" | ".join(filter(None, row_text)))
                    if table_text:
                        text.append("\n".join(table_text))
            
            return "\n\n".join(text)
    except Exception as e:
        raise Exception(f"Error processing PDF file: {str(e)}")

def extract_text_from_doc(file):
    """Extract text from a DOC/DOCX file"""
    # Save the uploaded file object to a BytesIO object
    doc_bytes = BytesIO(file.read())
    
    try:
        # Open the document using python-docx
        doc = DocxDocument(doc_bytes)
        
        # Extract text from paragraphs
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Only add non-empty cells
                        row_text.append(cell.text.strip())
                if row_text:  # Only add non-empty rows
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
        
    except Exception as e:
        raise Exception(f"Error processing DOC/DOCX file: {str(e)}") 

@bp.route('/project/<int:project_id>/update_notes', methods=['POST'])
@login_required
def update_notes(project_id):
    project = Project.query.get_or_404(project_id)
    
    # Get the notes from the form
    notes = request.form.get('notes', '').strip()
    
    # Update the project
    project.notes = notes
    db.session.commit()
    
    # Return success
    return jsonify({'success': True, 'message': 'Notes updated successfully'})

@bp.route('/project/<int:project_id>/update', methods=['POST'])
@login_required
def update_project(project_id):
    project = Project.query.get_or_404(project_id)
    
    try:
        # Update existing fields
        project.name_of_round = float(request.form.get('name_of_round'))
        project.file_number_db = request.form.get('file_number_db')
        project.scope = request.form.get('scope')
        project.region = request.form.get('region')
        project.countries_covered = request.form.get('countries_covered')
        project.integrity_partner_name = request.form.get('integrity_partner_name')
        project.partner_type = request.form.get('partner_type')
        project.project_partners = request.form.get('project_partners')
        project.wb_or_eib = request.form.get('wb_or_eib')
        
        # Add missing fields
        project.key_project_objectives = request.form.get('key_project_objectives')
        project.sectoral_scope = request.form.get('sectoral_scope')
        project.specific_sector = request.form.get('specific_sector')
        project.funding_amount_usd = float(request.form.get('funding_amount_usd')) if request.form.get('funding_amount_usd') else None
        project.duration = request.form.get('duration')
        project.start_year = int(request.form.get('start_year')) if request.form.get('start_year') else None
        project.end_year = int(request.form.get('end_year')) if request.form.get('end_year') else None
        project.wb_income_classification = request.form.get('wb_income_classification')
        project.corruption_quintile = request.form.get('corruption_quintile')
        project.cci = float(request.form.get('cci')) if request.form.get('cci') else None
        project.government_type_eiu = request.form.get('government_type_eiu')
        project.government_score_eiu = float(request.form.get('government_score_eiu')) if request.form.get('government_score_eiu') else None
        
        # Handle checkboxes
        project.final_external_evaluation = 'final_external_evaluation' in request.form
        project.final_report = 'final_report' in request.form
        project.full_proposal = 'full_proposal' in request.form
        project.workplan = 'workplan' in request.form
        project.baseline_assessment = 'baseline_assessment' in request.form
        
        # Handle optional other field
        project.other = request.form.get('other')

        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}) 

@bp.route('/test-prompt', methods=['GET', 'POST'])
@login_required
def test_prompt():
    response_text = None
    if request.method == 'POST':
        user_prompt = request.form.get('prompt', '')
        if user_prompt:
            try:
                client = OpenAI(api_key=Config.OPENAI_API_KEY)
                full_prompt = Config.DEFAULT_SYSTEM_PROMPT + "\n\n" + user_prompt
                
                response = client.chat.completions.create(
                    model="o3-mini",
                    messages=[
                        {"role": "user", "content": full_prompt}
                    ]
                )
                
                response_text = response.choices[0].message.content
                
            except Exception as e:
                flash(f'Error: {str(e)}', 'error')
    
    return render_template('main/test_prompt.html', 
                         response=response_text,
                         title='Test Prompt') 

@bp.route('/mark_response_reviewed/<int:response_id>', methods=['POST'])
@login_required
def mark_response_reviewed(response_id):
    response = EvaluationResponse.query.get_or_404(response_id)
    
    try:
        response.reviewed = True
        response.reviewed_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500 

@bp.route('/update_response/<int:response_id>', methods=['POST'])
@login_required
def update_response(response_id):
    response = EvaluationResponse.query.get_or_404(response_id)
    
    try:
        data = request.get_json()
        # Use the property setter for encryption
        response.response_text = data.get('response_text')
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500 

@bp.route('/project/<int:project_id>/update_metrics', methods=['POST'])
@login_required
def update_project_metrics(project_id):
    project = Project.query.get_or_404(project_id)
    
    try:
        # Update numeric counts - using type=float to allow decimal values
        project.num_pubpri_dialogues = request.form.get('num_pubpri_dialogues', type=float)
        project.num_legal_contribuntions = request.form.get('num_legal_contribuntions', type=float)
        project.num_implement_mechanisms = request.form.get('num_implement_mechanisms', type=float)
        project.num_voluntary_standards = request.form.get('num_voluntary_standards', type=float)
        project.num_voluntary_signatories = request.form.get('num_voluntary_signatories', type=float)
        project.num_organizations_supported = request.form.get('num_organizations_supported', type=float)
        project.num_new_courses = request.form.get('num_new_courses', type=float)
        project.num_individ_trained = request.form.get('num_individ_trained', type=float)
        project.num_training_activities = request.form.get('num_training_activities', type=float)
        project.num_organizaed_events = request.form.get('num_organizaed_events', type=float)
        project.num_event_attendees = request.form.get('num_event_attendees', type=float)
        project.num_publications = request.form.get('num_publications', type=float)
        
        # Update ratings
        project.rate_output_achieved = request.form.get('rate_output_achieved', type=float)
        project.rate_impact_evidence = request.form.get('rate_impact_evidence')  # String value
        project.rate_sustainability = request.form.get('rate_sustainability', type=float)
        project.rate_project_design = request.form.get('rate_project_design', type=float)
        project.rate_project_management = request.form.get('rate_project_management')  # String value
        project.rate_quality_evaluation = request.form.get('rate_quality_evaluation')  # String value
        project.rate_impact_progress = request.form.get('rate_impact_progress', type=float)
        
        # Update significance ratings
        project.rate_signif_frameworks = request.form.get('rate_signif_frameworks')
        project.rate_signif_practices = request.form.get('rate_signif_practices')
        project.rate_signif_capacity = request.form.get('rate_signif_capacity')
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}) 

@bp.route('/projects/document/<int:document_id>/preview')
@login_required
def preview_document(document_id):
    try:
        document = Document.query.get_or_404(document_id)
        
        # Check if user has access to this document's project
        if not current_user.is_admin:
            project = Project.query.get(document.project_id)
            if not project or not project.active:
                return jsonify({
                    'success': False,
                    'error': 'Access denied'
                }), 403
        
        # Get decrypted content
        content = document.get_content()
        
        return jsonify({
            'success': True,
            'content': content,
            'filename': document.filename
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500 

@bp.route('/project_table')
@login_required
def project_table():
    # Update query to sort by Round and File # in ascending order
    projects = Project.query.order_by(
        Project.name_of_round.asc(),
        Project.file_number_db.asc()
    ).all()
    return render_template('main/project_table.html', projects=projects) 

@bp.route('/interviews')
@login_required
def interviews():
    # Use _content in the query to match the actual column name
    interviews = db.session.query(Interview).order_by(Interview.uploaded_at.desc()).all()
    return render_template('main/interviews.html', interviews=interviews) 

@bp.route('/upload_interview', methods=['GET', 'POST'])
@login_required
def upload_interview():
    if request.method == 'POST':
        # Check if a file was uploaded
        if 'file' not in request.files:
            flash('No file selected', 'danger')
            return redirect(request.url)
            
        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'danger')
            return redirect(request.url)

        if file:
            # Determine the file type and extract content accordingly
            content = ""
            if file.filename.endswith('.docx'):
                # Read content from Word document
                doc = DocxDocument(file)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                # Read content from text file
                content = file.read().decode('utf-8')

            # Calculate word count
            word_count = len(content.split())

            # Create new interview
            interview = Interview(
                name=request.form.get('name', file.filename),
                filename=file.filename,
                content=content,
                word_count=word_count,  # Save the word count
                uploaded_by=current_user
            )
            
            db.session.add(interview)
            db.session.commit()
            
            flash('Interview uploaded successfully', 'success')
            return redirect(url_for('main.interviews'))

    return render_template('main/upload_interview.html') 

@bp.route('/interview/<int:interview_id>')
@login_required
def view_interview(interview_id):
    interview = Interview.query.get_or_404(interview_id)
    return render_template('main/view_interview.html', interview=interview)

@bp.route('/interview/<int:interview_id>/delete', methods=['POST'])
@login_required
def delete_interview(interview_id):
    interview = Interview.query.get_or_404(interview_id)
    db.session.delete(interview)
    db.session.commit()
    flash('Interview deleted successfully', 'success')
    return redirect(url_for('main.interviews')) 

@bp.route('/interview_questions')
@login_required
def interview_questions():
    questions = InterviewQuestion.query.all()
    return render_template('main/interview_questions.html', questions=questions)

@bp.route('/edit_interview_question/<int:question_id>', methods=['GET', 'POST'])
@login_required
def edit_interview_question(question_id):
    question = InterviewQuestion.query.get_or_404(question_id)
    
    if request.method == 'POST':
        question.title = request.form['title']
        question.text = request.form['text']
        db.session.commit()
        flash('Interview question updated successfully', 'success')
        return redirect(url_for('main.interview_questions'))
    
    return render_template('main/edit_interview_question.html', question=question)

@bp.route('/submit_interview_question/<int:question_id>', methods=['GET', 'POST'])
@login_required
def submit_interview_question(question_id=None):
    question_text = ""
    
    if question_id is not None:
        # Load the interview question text by ID
        question = InterviewQuestion.query.get_or_404(question_id)
        question_text = question.text  # Load the text of the question

    if request.method == 'POST':
        question_text = request.form['question_text']
        system_prompt = request.form['system_prompt']  # Get the system prompt from the form
        
        # Fetch all interview texts
        interviews = Interview.query.all()
        interview_texts = [interview.content for interview in interviews if interview.content]  # Collect texts
        
        # Combine all interview texts into a single string
        combined_interview_texts = "\n".join(interview_texts)

        # Send the question, system prompt, and interview texts to the local LLM
        response = requests.post(
            'http://127.0.0.1:1234/v1/chat/completions',
            json={
                "messages": [
                    {"role": "system", "content": system_prompt},  # System prompt
                    {"role": "user", "content": question_text + "\n\n" + combined_interview_texts}  # User question with context
                ],
                "model": "llama3"  # This is just a placeholder - LM Studio generally ignores this field
            }
        )
        
        if response.status_code == 200:
            llm_response = response.json()
            llm_content = llm_response['choices'][0]['message']['content']
            llm_content = llm_content.replace('\n', '<br>')  # Convert line returns to <br> for HTML display
            flash('Question submitted successfully!', 'success')
            return render_template('main/submit_interview_question.html', llm_response=llm_content, question_text=question_text)
        else:
            flash('Error submitting question to LLM: ' + response.text, 'danger')
    
    return render_template('main/submit_interview_question.html', question_text=question_text) 

def get_projects_with_all_responses():
    """
    Count projects that have responses GENERATED for all questions.
    """
    # Get all questions
    all_questions = ProjectQuestion.query.all()
    total_questions = len(all_questions)
    question_ids = [q.id for q in all_questions]
    
    if total_questions == 0:
        return 0
    
    # Get all projects
    all_projects = Project.query.all()
    projects_with_all_responses = 0
    
    for project in all_projects:
        # Get all responses for this project
        responses = EvaluationResponse.query.filter_by(project_id=project.id).all()
        
        # Filter valid responses in Python instead of SQL
        valid_responses = [
            response for response in responses 
            if response.response_text and 
            response.response_text.strip() and 
            response.response_text != "Processing..."
        ]
        
        # Get unique question IDs that have responses
        answered_question_ids = set(response.question_id for response in valid_responses)
        
        # Check if all questions have been answered
        if set(question_ids).issubset(answered_question_ids):
            projects_with_all_responses += 1
    
    return projects_with_all_responses

def get_projects_with_all_reviewed():
    """
    Count projects that have ALL responses REVIEWED.
    """
    # Get all questions
    all_questions = ProjectQuestion.query.all()
    total_questions = len(all_questions)
    question_ids = [q.id for q in all_questions]
    
    if total_questions == 0:
        return 0
    
    # Get all projects
    all_projects = Project.query.all()
    projects_with_all_reviewed = 0
    
    for project in all_projects:
        # Get all reviewed responses for this project
        responses = EvaluationResponse.query.filter_by(
            project_id=project.id,
            reviewed=True
        ).all()
        
        # Filter valid responses in Python
        valid_responses = [
            response for response in responses 
            if response.response_text and 
            response.response_text.strip() and 
            response.response_text != "Processing..."
        ]
        
        # Get unique question IDs that have reviewed responses
        reviewed_question_ids = set(response.question_id for response in valid_responses)
        
        # Check if all questions have been reviewed
        if set(question_ids).issubset(reviewed_question_ids):
            projects_with_all_reviewed += 1
    
    return projects_with_all_reviewed

@bp.route('/export_projects/<format>')
@login_required
def export_projects(format):
    # Apply the same sorting as the table view
    projects = Project.query.order_by(
        Project.name_of_round.asc(),
        Project.file_number_db.asc()
    ).all()
    
    if format == 'csv':
        # Create a CSV in memory
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header with all columns from the table
        writer.writerow([
            'Partner Name', 'Round', 'File #', 'Partner Type', 'Project Partners', 'WB/EIB',
            'Scope', 'Region', 'Countries', 'Sectoral Scope', 'Specific Sector', 'Funding (USD M)',
            'Duration', 'Start Year', 'End Year', 'WB Income Class', 'Corruption Quintile', 'CCI',
            'Gov Type (EIU)', 'Gov Score (EIU)', 'Key Objectives', 'Public-Private Dialogues',
            'Legal Contributions', 'Implementation Mechanisms', 'Voluntary Standards', 
            'Voluntary Signatories', 'Organizations Supported', 'New Courses', 'Individuals Trained',
            'Training Activities', 'Events Organized', 'Event Attendees', 'Publications',
            'Output Achievement', 'Impact Evidence', 'Sustainability', 'Project Design',
            'Project Management', 'Evaluation Quality', 'Impact Progress', 'Framework Significance',
            'Practice Significance', 'Capacity Significance'
        ])
        
        # Write data for all columns
        for project in projects:
            writer.writerow([
                project.integrity_partner_name,
                project.name_of_round,
                project.file_number_db,
                project.partner_type,
                project.project_partners,
                project.wb_or_eib,
                project.scope,
                project.region,
                project.countries_covered,
                project.sectoral_scope,
                project.specific_sector,
                project.funding_amount_usd,
                project.duration,
                project.start_year,
                project.end_year,
                project.wb_income_classification,
                project.corruption_quintile,
                project.cci,
                project.government_type_eiu,
                project.government_score_eiu,
                project.key_project_objectives,
                project.num_pubpri_dialogues,
                project.num_legal_contribuntions,
                project.num_implement_mechanisms,
                project.num_voluntary_standards,
                project.num_voluntary_signatories,
                project.num_organizations_supported,
                project.num_new_courses,
                project.num_individ_trained,
                project.num_training_activities,
                project.num_organizaed_events,
                project.num_event_attendees,
                project.num_publications,
                project.rate_output_achieved,
                project.rate_impact_evidence,
                project.rate_sustainability,
                project.rate_project_design,
                project.rate_project_management,
                project.rate_quality_evaluation,
                project.rate_impact_progress,
                project.rate_signif_frameworks,
                project.rate_signif_practices,
                project.rate_signif_capacity
            ])
        
        # Prepare response
        output.seek(0)
        return send_file(
            io.BytesIO(output.getvalue().encode('utf-8')),
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'projects_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        )
    
    elif format == 'excel':
        try:
            # Create an Excel file in memory
            output = io.BytesIO()
            workbook = xlsxwriter.Workbook(output)
            worksheet = workbook.add_worksheet()
            
            # Add header with formatting
            header_format = workbook.add_format({'bold': True, 'bg_color': '#4e73df', 'font_color': 'white'})
            headers = [
                'Partner Name', 'Round', 'File #', 'Partner Type', 'Project Partners', 'WB/EIB',
                'Scope', 'Region', 'Countries', 'Sectoral Scope', 'Specific Sector', 'Funding (USD M)',
                'Duration', 'Start Year', 'End Year', 'WB Income Class', 'Corruption Quintile', 'CCI',
                'Gov Type (EIU)', 'Gov Score (EIU)', 'Key Objectives', 'Public-Private Dialogues',
                'Legal Contributions', 'Implementation Mechanisms', 'Voluntary Standards', 
                'Voluntary Signatories', 'Organizations Supported', 'New Courses', 'Individuals Trained',
                'Training Activities', 'Events Organized', 'Event Attendees', 'Publications',
                'Output Achievement', 'Impact Evidence', 'Sustainability', 'Project Design',
                'Project Management', 'Evaluation Quality', 'Impact Progress', 'Framework Significance',
                'Practice Significance', 'Capacity Significance'
            ]
            
            for col, header in enumerate(headers):
                worksheet.write(0, col, header, header_format)
            
            # Write data for all columns
            for row, project in enumerate(projects, start=1):
                data = [
                    project.integrity_partner_name,
                    project.name_of_round,
                    project.file_number_db,
                    project.partner_type,
                    project.project_partners,
                    project.wb_or_eib,
                    project.scope,
                    project.region,
                    project.countries_covered,
                    project.sectoral_scope,
                    project.specific_sector,
                    project.funding_amount_usd,
                    project.duration,
                    project.start_year,
                    project.end_year,
                    project.wb_income_classification,
                    project.corruption_quintile,
                    project.cci,
                    project.government_type_eiu,
                    project.government_score_eiu,
                    project.key_project_objectives,
                    project.num_pubpri_dialogues,
                    project.num_legal_contribuntions,
                    project.num_implement_mechanisms,
                    project.num_voluntary_standards,
                    project.num_voluntary_signatories,
                    project.num_organizations_supported,
                    project.num_new_courses,
                    project.num_individ_trained,
                    project.num_training_activities,
                    project.num_organizaed_events,
                    project.num_event_attendees,
                    project.num_publications,
                    project.rate_output_achieved,
                    project.rate_impact_evidence,
                    project.rate_sustainability,
                    project.rate_project_design,
                    project.rate_project_management,
                    project.rate_quality_evaluation,
                    project.rate_impact_progress,
                    project.rate_signif_frameworks,
                    project.rate_signif_practices,
                    project.rate_signif_capacity
                ]
                
                for col, value in enumerate(data):
                    worksheet.write(row, col, value)
            
            # Auto-adjust column widths
            for col, header in enumerate(headers):
                worksheet.set_column(col, col, len(header) + 5)
            
            # Add text wrapping for columns with potentially long content
            wrap_format = workbook.add_format({'text_wrap': True})
            for col in [4, 9, 10, 20]:  # Project Partners, Sectoral Scope, Specific Sector, Key Objectives
                worksheet.set_column(col, col, 30, wrap_format)
            
            workbook.close()
            
            # Prepare response
            output.seek(0)
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'projects_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
        except NameError:
            flash('Excel export is not available. Please install xlsxwriter package.', 'warning')
            return redirect(url_for('main.project_table'))
    
    # Default fallback
    flash('Invalid export format specified.', 'danger')
    return redirect(url_for('main.project_table'))

